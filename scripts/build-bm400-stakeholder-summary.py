from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

OUT = Path("outputs/bm400-framework-guide")
OUT.mkdir(parents=True, exist_ok=True)

mapping = [
    ("Religius", "Akhlak", "Karakter", "Akhlak mulia dalam interaksi sehari-hari", "Akhlak Islami dalam Keseharian", "4"),
    ("Religius", "Ibadah", "Karakter", "Konsistensi dan adab dalam beribadah", "Kedisiplinan dan Kekhusyukan Ibadah", "4"),
    ("Religius", "Praktik Baik", "Karakter", "Kepedulian dan amal kebajikan secara sukarela", "Kebiasaan Berbuat Baik", "4"),
    ("Religius", "Muamalah", "Soft Skill", "Muamalah: membangun hubungan sosial yang harmonis", "Interaksi Sosial yang Sehat", "4"),
    ("Nasionalis", "Patriotisme", "Karakter", "Kecintaan dan kepedulian terhadap bangsa", "Patriotisme dalam Keseharian", "4"),
    ("Nasionalis", "Toleransi", "Soft Skill", "Toleransi terhadap keberagaman", "Sikap Toleran", "3"),
    ("Nasionalis", "Demokrasi", "Soft Skill", "Demokrasi dan musyawarah", "Kecakapan Bermusyawarah", "3"),
    ("Internasionalis", "Berwawasan Global", "Mental", "Berwawasan global dan terbuka terhadap keberagaman", "Keterbukaan Wawasan Global", "3"),
    ("Internasionalis", "Unggul & Kompetitif", "Mental", "Semangat unggul dan kompetitif secara sehat", "Orientasi Keunggulan", "3"),
    ("Internasionalis", "Kemampuan Adaptasi", "Mental", "Kemampuan beradaptasi terhadap perubahan", "Adaptasi Lintas Lingkungan", "3"),
    ("Internasionalis", "Kemandirian Berpikir", "Mental", "Kemandirian berpikir dan pengambilan keputusan", "Berpikir Kritis dan Mandiri", "3"),
    ("Internasionalis", "Kolaborasi Lintas Budaya", "Soft Skill", "Kolaborasi lintas budaya", "Kerja Sama Lintas Budaya", "3"),
    ("Internasionalis", "Integritas Global", "Karakter", "Integritas dalam lingkungan global", "Integritas dan Etika Global", "3"),
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def cell_text(cell, text, bold=False, color="1F2937", size=8.3):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text)); run.bold = bold; run.font.name = "Calibri"; run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table_doc(doc, rows, widths, size=8.3):
    t = doc.add_table(rows=0, cols=len(widths)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i]); cell_text(cells[i], value, ri == 0, "FFFFFF" if ri == 0 else "1F2937", 8 if ri == 0 else size)
            shade(cells[i], "1F4D78" if ri == 0 else ("F4F7FA" if ri % 2 == 0 else "FFFFFF"))
    return t

def build_docx(path):
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = Inches(.75); sec.bottom_margin = Inches(.7); sec.left_margin = Inches(.8); sec.right_margin = Inches(.8)
    sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)
    normal = doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor(31,41,55); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.08
    for name, size, color, before, after in [("Heading 1",16,"1F4D78",14,7),("Heading 2",12.5,"2E74B5",10,5)]:
        s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    h=sec.header.paragraphs[0]; h.text="Framework Tambahan BM400"; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=RGBColor(107,114,128)
    f=sec.footer.paragraphs[0]; f.text="SMA Bakti Mulya 400"; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.runs[0].font.size=Pt(8); f.runs[0].font.color.rgb=RGBColor(107,114,128)
    p=doc.add_paragraph(); r=p.add_run("FRAMEWORK TAMBAHAN BM400"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(11,37,69)
    p=doc.add_paragraph("Penjelasan konversi Naskah Akademis Tiga Pilar"); p.runs[0].font.size=Pt(12); p.runs[0].font.color.rgb=RGBColor(75,85,99)
    doc.add_heading("Ringkasan Eksekutif",1)
    doc.add_paragraph("Framework BM400 bukan pengganti framework CDS yang sudah ada. Ia adalah lapisan tambahan khusus untuk SMA Bakti Mulya 400 yang menambahkan 13 dimensi dari Naskah Akademis Tiga Pilar: Religius, Nasionalis, dan Internasionalis.")
    doc.add_paragraph("Konversi dilakukan agar dokumen kebijakan yang bersifat naratif dapat dipakai secara konsisten oleh CMS: setiap dimensi dijadikan satu tema, setiap tema memiliki satu indikator operasional, dan setiap indikator memiliki 3-4 sub-indikator yang dapat diamati dalam percakapan asesmen.")
    table_doc(doc, [["Lapisan","Sebelum","Tambahan BM400","Total BM400"],["Karakter","40 tema","+5 tema","45 tema"],["Mental","34 tema","+4 tema","38 tema"],["Soft Skill","14 tema","+4 tema","18 tema"],["Total dimensi/tema tambahan","-","13","13"]],[2.2,1.2,1.3,1.5])
    doc.add_heading("1. Prinsip Konversi",1)
    doc.add_paragraph("Dokumen sumber memiliki struktur pilar -> dimensi -> uraian operasional -> contoh kejadian/kompetensi. Struktur tersebut dipertahankan, lalu diterjemahkan ke struktur CMS yang sudah dipakai aplikasi.")
    table_doc(doc, [["Struktur sumber PDF","Struktur CMS/AI","Fungsi"],["Pilar","Kategori","Karakter, Mental, atau Soft Skill"],["Dimensi","Tema","Unit utama yang dilacak dalam perkembangan siswa"],["Uraian dimensi","Penjelasan tema","Konteks makna dan alasan pentingnya tema"],["Kompetensi/contoh kejadian","Sub-indikator","Bukti perilaku yang dapat ditandai terpenuhi atau menurun"]],[2.2,2.0,2.5])
    doc.add_paragraph("Pengelompokan kategori menggunakan sifat utama kompetensinya: nilai dan perilaku moral masuk Karakter; pola pikir, daya adaptasi, dan orientasi belajar masuk Mental; kemampuan berhubungan dan bekerja sama masuk Soft Skill.", style="Intense Quote")
    doc.add_heading("2. Pemetaan 13 Dimensi",1)
    doc.add_paragraph("Tabel berikut menunjukkan jejak konversi dari dimensi dalam dokumen sumber sampai ke tema dan indikator.")
    table_doc(doc, [["Pilar","Dimensi sumber","Kategori CMS","Tema hasil konversi","Indikator","Sub"]]+mapping,[.85,1.15,.8,2.05,1.75,.4],8.1)
    doc.add_heading("3. Contoh Konversi Lengkap",1)
    doc.add_heading("Contoh A - Religius / Akhlak",2)
    doc.add_paragraph("Dimensi Akhlak tidak dimasukkan sebagai kalimat abstrak saja. Ia dioperasionalkan menjadi bukti perilaku yang dapat muncul dalam wawancara, misalnya pengembalian barang, pengakuan kesalahan, kesantunan, dan tanggung jawab.")
    table_doc(doc, [["Level","Hasil"],["Kategori","Karakter"],["Tema","Akhlak mulia dalam interaksi sehari-hari"],["Indikator","Akhlak Islami dalam Keseharian"],["Sub-indikator","Mengembalikan barang yang ditemukan kepada pemiliknya; Mengakui kesalahan yang diperbuat; Berbicara santun kepada guru dan teman; Menepati janji dan menyelesaikan tugas dengan penuh tanggung jawab"]],[1.35,5.35])
    doc.add_heading("Contoh B - Internasionalis / Berwawasan Global",2)
    doc.add_paragraph("Dimensi Berwawasan Global ditempatkan di Mental karena fokusnya adalah keterbukaan cara pandang, pemahaman isu, dan kebiasaan mencari pengetahuan. Ia tidak dikategorikan sebagai Soft Skill meskipun dapat berdampak pada interaksi sosial.")
    table_doc(doc, [["Level","Hasil"],["Kategori","Mental"],["Tema","Berwawasan global dan terbuka terhadap keberagaman"],["Indikator","Keterbukaan Wawasan Global"],["Sub-indikator","Menunjukkan keterbukaan terhadap budaya, bahasa, dan cara pandang yang berbeda; Memahami isu-isu global secara objektif; Mengikuti perkembangan ilmu pengetahuan dan isu global dari sumber terpercaya"]],[1.35,5.35])
    doc.save(path)

def build_md(path):
    lines=["# Framework Tambahan BM400","","## Ringkasan Eksekutif","","Framework BM400 bukan pengganti framework CDS yang sudah ada. Ia adalah lapisan tambahan khusus untuk SMA Bakti Mulya 400 yang menambahkan 13 dimensi dari Naskah Akademis Tiga Pilar: Religius, Nasionalis, dan Internasionalis.","","Konversi dilakukan agar dokumen kebijakan yang bersifat naratif dapat dipakai secara konsisten oleh CMS: setiap dimensi dijadikan satu tema, setiap tema memiliki satu indikator operasional, dan setiap indikator memiliki 3-4 sub-indikator yang dapat diamati dalam percakapan asesmen.","","| Lapisan | Sebelum | Tambahan BM400 | Total BM400 |","|---|---:|---:|---:|","| Karakter | 40 tema | +5 tema | 45 tema |","| Mental | 34 tema | +4 tema | 38 tema |","| Soft Skill | 14 tema | +4 tema | 18 tema |","| Total dimensi/tema tambahan | - | 13 | 13 |","","## 1. Prinsip Konversi","","Dokumen sumber memiliki struktur pilar -> dimensi -> uraian operasional -> contoh kejadian/kompetensi. Struktur tersebut dipertahankan, lalu diterjemahkan ke struktur CMS yang sudah dipakai aplikasi.","","| Struktur sumber PDF | Struktur CMS/AI | Fungsi |","|---|---|---|","| Pilar | Kategori | Karakter, Mental, atau Soft Skill |","| Dimensi | Tema | Unit utama yang dilacak dalam perkembangan siswa |","| Uraian dimensi | Penjelasan tema | Konteks makna dan alasan pentingnya tema |","| Kompetensi/contoh kejadian | Sub-indikator | Bukti perilaku yang dapat ditandai terpenuhi atau menurun |","","Pengelompokan kategori menggunakan sifat utama kompetensinya: nilai dan perilaku moral masuk Karakter; pola pikir, daya adaptasi, dan orientasi belajar masuk Mental; kemampuan berhubungan dan bekerja sama masuk Soft Skill.","","## 2. Pemetaan 13 Dimensi","","Tabel berikut menunjukkan jejak konversi dari dimensi dalam dokumen sumber sampai ke tema dan indikator.","","| Pilar | Dimensi sumber | Kategori CMS | Tema hasil konversi | Indikator | Sub |","|---|---|---|---|---|---:|"]
    lines += ["| "+" | ".join(r)+" |" for r in mapping]
    lines += ["","## 3. Contoh Konversi Lengkap","","### Contoh A - Religius / Akhlak","","Dimensi Akhlak tidak dimasukkan sebagai kalimat abstrak saja. Ia dioperasionalkan menjadi bukti perilaku yang dapat muncul dalam wawancara, misalnya pengembalian barang, pengakuan kesalahan, kesantunan, dan tanggung jawab.","","| Level | Hasil |","|---|---|","| Kategori | Karakter |","| Tema | Akhlak mulia dalam interaksi sehari-hari |","| Indikator | Akhlak Islami dalam Keseharian |","| Sub-indikator | Mengembalikan barang yang ditemukan kepada pemiliknya; Mengakui kesalahan yang diperbuat; Berbicara santun kepada guru dan teman; Menepati janji dan menyelesaikan tugas dengan penuh tanggung jawab |","","### Contoh B - Internasionalis / Berwawasan Global","","Dimensi Berwawasan Global ditempatkan di Mental karena fokusnya adalah keterbukaan cara pandang, pemahaman isu, dan kebiasaan mencari pengetahuan. Ia tidak dikategorikan sebagai Soft Skill meskipun dapat berdampak pada interaksi sosial.","","| Level | Hasil |","|---|---|","| Kategori | Mental |","| Tema | Berwawasan global dan terbuka terhadap keberagaman |","| Indikator | Keterbukaan Wawasan Global |","| Sub-indikator | Menunjukkan keterbukaan terhadap budaya, bahasa, dan cara pandang yang berbeda; Memahami isu-isu global secara objektif; Mengikuti perkembangan ilmu pengetahuan dan isu global dari sumber terpercaya |",""]
    path.write_text("\n".join(lines),encoding="utf-8")

def build_pdf(path):
    ss=getSampleStyleSheet(); ss.add(ParagraphStyle(name="T",parent=ss["Title"],fontName="Helvetica-Bold",fontSize=19,leading=23,textColor=colors.HexColor("#0B2545"),spaceAfter=5)); ss.add(ParagraphStyle(name="H1b",parent=ss["Heading1"],fontName="Helvetica-Bold",fontSize=14,leading=17,textColor=colors.HexColor("#1F4D78"),spaceBefore=10,spaceAfter=5)); ss.add(ParagraphStyle(name="H2b",parent=ss["Heading2"],fontName="Helvetica-Bold",fontSize=11,leading=14,textColor=colors.HexColor("#2E74B5"),spaceBefore=8,spaceAfter=4)); ss.add(ParagraphStyle(name="B",parent=ss["BodyText"],fontName="Helvetica",fontSize=9,leading=12.5,textColor=colors.HexColor("#1F2937"),spaceAfter=5)); ss.add(ParagraphStyle(name="S",parent=ss["BodyText"],fontName="Helvetica",fontSize=7.1,leading=8.5,textColor=colors.HexColor("#1F2937")))
    doc=SimpleDocTemplate(str(path),pagesize=A4,leftMargin=16*mm,rightMargin=16*mm,topMargin=16*mm,bottomMargin=15*mm); story=[]; P=lambda x,s="B":Paragraph(x,ss[s])
    story += [P("FRAMEWORK TAMBAHAN BM400","T"),P("Penjelasan konversi Naskah Akademis Tiga Pilar")]
    story += [P("Ringkasan Eksekutif","H1b"),P("Framework BM400 bukan pengganti framework CDS yang sudah ada. Ia adalah lapisan tambahan khusus untuk SMA Bakti Mulya 400 yang menambahkan 13 dimensi dari Naskah Akademis Tiga Pilar: Religius, Nasionalis, dan Internasionalis."),P("Konversi dilakukan agar dokumen kebijakan yang bersifat naratif dapat dipakai secara konsisten oleh CMS: setiap dimensi dijadikan satu tema, setiap tema memiliki satu indikator operasional, dan setiap indikator memiliki 3-4 sub-indikator yang dapat diamati dalam percakapan asesmen.")]
    def tbl(data,widths,small=False):
        d=[[P(str(v),"S" if small else "B") for v in row] for row in data]; t=Table(d,colWidths=[w*mm for w in widths],repeatRows=1); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4D78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.25,colors.HexColor("#CBD5E1")),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F4F7FA")]),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); return t
    story += [tbl([["Lapisan","Sebelum","Tambahan BM400","Total BM400"],["Karakter","40 tema","+5 tema","45 tema"],["Mental","34 tema","+4 tema","38 tema"],["Soft Skill","14 tema","+4 tema","18 tema"],["Total dimensi/tema tambahan","-","13","13"]],[52,30,40,42])]
    story += [P("1. Prinsip Konversi","H1b"),P("Dokumen sumber memiliki struktur pilar -> dimensi -> uraian operasional -> contoh kejadian/kompetensi. Struktur tersebut dipertahankan, lalu diterjemahkan ke struktur CMS yang sudah dipakai aplikasi."),tbl([["Struktur sumber PDF","Struktur CMS/AI","Fungsi"],["Pilar","Kategori","Karakter, Mental, atau Soft Skill"],["Dimensi","Tema","Unit utama yang dilacak dalam perkembangan siswa"],["Uraian dimensi","Penjelasan tema","Konteks makna dan alasan pentingnya tema"],["Kompetensi/contoh kejadian","Sub-indikator","Bukti perilaku yang dapat ditandai terpenuhi atau menurun"]],[48,43,73],True),P("Pengelompokan kategori menggunakan sifat utama kompetensinya: nilai dan perilaku moral masuk Karakter; pola pikir, daya adaptasi, dan orientasi belajar masuk Mental; kemampuan berhubungan dan bekerja sama masuk Soft Skill.")]
    story += [P("2. Pemetaan 13 Dimensi","H1b"),P("Tabel berikut menunjukkan jejak konversi dari dimensi dalam dokumen sumber sampai ke tema dan indikator."),tbl([["Pilar","Dimensi sumber","Kategori CMS","Tema hasil konversi","Indikator","Sub"]]+mapping,[22,27,22,48,43,10],True)]
    story += [P("3. Contoh Konversi Lengkap","H1b"),P("Contoh A - Religius / Akhlak","H2b"),P("Dimensi Akhlak tidak dimasukkan sebagai kalimat abstrak saja. Ia dioperasionalkan menjadi bukti perilaku yang dapat muncul dalam wawancara, misalnya pengembalian barang, pengakuan kesalahan, kesantunan, dan tanggung jawab."),tbl([["Level","Hasil"],["Kategori","Karakter"],["Tema","Akhlak mulia dalam interaksi sehari-hari"],["Indikator","Akhlak Islami dalam Keseharian"],["Sub-indikator","Mengembalikan barang yang ditemukan kepada pemiliknya; Mengakui kesalahan yang diperbuat; Berbicara santun kepada guru dan teman; Menepati janji dan menyelesaikan tugas dengan penuh tanggung jawab"]],[38,132],True),P("Contoh B - Internasionalis / Berwawasan Global","H2b"),P("Dimensi Berwawasan Global ditempatkan di Mental karena fokusnya adalah keterbukaan cara pandang, pemahaman isu, dan kebiasaan mencari pengetahuan. Ia tidak dikategorikan sebagai Soft Skill meskipun dapat berdampak pada interaksi sosial."),tbl([["Level","Hasil"],["Kategori","Mental"],["Tema","Berwawasan global dan terbuka terhadap keberagaman"],["Indikator","Keterbukaan Wawasan Global"],["Sub-indikator","Menunjukkan keterbukaan terhadap budaya, bahasa, dan cara pandang yang berbeda; Memahami isu-isu global secara objektif; Mengikuti perkembangan ilmu pengetahuan dan isu global dari sumber terpercaya"]],[38,132],True)]
    def footer(c,d): c.saveState(); c.setFont("Helvetica",7); c.setFillColor(colors.HexColor("#6B7280")); c.drawCentredString(A4[0]/2,8*mm,f"Framework Tambahan BM400 | Halaman {d.page}"); c.restoreState()
    doc.build(story,onFirstPage=footer,onLaterPages=footer)

build_docx(OUT/"penjelasan-framework-bm400.docx")
build_md(OUT/"penjelasan-framework-bm400.md")
build_pdf(OUT/"penjelasan-framework-bm400.pdf")
print("Regenerated simplified stakeholder documents")
