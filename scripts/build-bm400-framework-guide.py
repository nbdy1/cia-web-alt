from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

OUT = Path("outputs/bm400-framework-guide")
OUT.mkdir(parents=True, exist_ok=True)

ORG_ID = "0bc3db16-d270-42d9-893a-c233a6b83800"
rows = [
    ("Religius", "Akhlak", "Karakter", "Akhlak mulia dalam interaksi sehari-hari", "Akhlak Islami dalam Keseharian", "4"),
    ("Religius", "Ibadah", "Karakter", "Konsistensi dan adab dalam beribadah", "Kedisiplinan dan Kekhusyukan Ibadah", "4"),
    ("Religius", "Praktik Baik", "Karakter", "Kepedulian dan amal kebajikan secara sukarela", "Kebiasaan Berbuat Baik", "4"),
    ("Religius", "Muamalah", "Soft Skill", "Muamalah: membangun hubungan sosial yang harmonis", "Interaksi Sosial yang Sehat", "4"),
    ("Nasionalis", "Patriotisme", "Karakter", "Kecintaan dan kepedulian terhadap bangsa", "Patriotisme dalam Keseharian", "4"),
    ("Nasionalis", "Toleransi", "Soft Skill", "Toleransi terhadap keberagaman", "Sikap Toleran", "3"),
    ("Nasionalis", "Demokrasi", "Soft Skill", "Demokrasi dan musyawarah", "Kecakapan Bermusyawarah", "3"),
    ("Internasionalis", "Berwawasan Global", "Mental", "Berwawasan global dan terbuka terhadap keberagaman", "Keterbukaan Wawasan Global", "3"),
    ("Internasionalis", "Unggul & Kompetitif", "Mental", "Semangat unggul dan kompetitif secara sehat", "Orientasi Keunggulan", "3"),
    ("Internasionalis", "Kemampuan Adaptasi", "Mental", "Kemampuan beradaptasi terhadap perubahan", "Adaptasi Lintas Lingkungan", "3"),
    ("Internasionalis", "Kemandirian Berpikir", "Mental", "Kemandirian berpikir dan pengambilan keputusan", "Berpikir Kritis dan Mandiri", "3"),
    ("Internasionalis", "Kolaborasi Lintas Budaya", "Soft Skill", "Kolaborasi lintas budaya", "Kerja Sama Lintas Budaya", "3"),
    ("Internasionalis", "Integritas Global", "Karakter", "Integritas dalam lingkungan global", "Integritas dan Etika Global", "3"),
]

details = {
    "Akhlak": ["Mengembalikan barang yang ditemukan kepada pemiliknya", "Mengakui kesalahan yang diperbuat", "Berbicara santun kepada guru dan teman", "Menepati janji dan menyelesaikan tugas dengan penuh tanggung jawab"],
    "Ibadah": ["Mengikuti salat berjamaah dengan tertib", "Berdoa sebelum dan sesudah kegiatan", "Membaca Al-Qur'an dengan sungguh-sungguh", "Menjaga ketenangan dan adab saat beribadah"],
    "Praktik Baik": ["Menjaga kebersihan kelas dan fasilitas sekolah", "Membantu guru atau teman tanpa diminta", "Mengikuti kegiatan bakti sosial atau menjadi relawan kegiatan sekolah", "Mengajak teman melakukan kegiatan positif"],
    "Muamalah": ["Membantu teman yang mengalami kesulitan", "Menghormati perbedaan pendapat", "Bekerja sama dalam kelompok dan menjaga amanah", "Menyelesaikan konflik secara santun"],
    "Patriotisme": ["Mengikuti upacara bendera dengan tertib dan menunjukkan sikap hormat", "Aktif dalam kegiatan kepramukaan atau peringatan hari besar nasional", "Berinisiatif menggalang bantuan untuk korban bencana", "Menunjukkan kepedulian terhadap kondisi bangsa dan lingkungan sekitar"],
    "Toleransi": ["Bekerja sama dalam kelompok dengan teman dari latar belakang berbeda", "Membela teman yang menjadi sasaran perundungan atas dasar perbedaan", "Menyampaikan pendapat berbeda dengan bahasa yang santun"],
    "Demokrasi": ["Mengajukan usulan secara terbuka dalam rapat kelas atau organisasi siswa", "Menerima hasil musyawarah meskipun berbeda dengan pendapat pribadi", "Menjalankan tanggung jawab kepemimpinan atau keanggotaan organisasi secara konsisten"],
    "Berwawasan Global": ["Menunjukkan keterbukaan terhadap budaya, bahasa, dan cara pandang yang berbeda", "Memahami isu-isu global secara objektif", "Mengikuti perkembangan ilmu pengetahuan dan isu global dari sumber terpercaya"],
    "Unggul & Kompetitif": ["Menunjukkan kerja keras dan disiplin untuk mencapai standar akademik terbaik", "Menunjukkan kemauan untuk terus berkembang", "Bersaing secara sehat di tingkat nasional maupun internasional"],
    "Kemampuan Adaptasi": ["Menyesuaikan diri dengan lingkungan akademik atau sosial yang baru", "Menyesuaikan diri terhadap perkembangan teknologi", "Menunjukkan kepercayaan diri berinteraksi di lingkungan multikultural"],
    "Kemandirian Berpikir": ["Menganalisis informasi secara kritis sebelum mengambil kesimpulan", "Mengambil keputusan berdasarkan data dan bertanggung jawab atasnya", "Menyusun perencanaan pendidikan atau portofolio secara mandiri"],
    "Kolaborasi Lintas Budaya": ["Bekerja sama secara efektif dengan individu dari berbagai latar belakang budaya", "Membangun komunikasi yang positif dalam lingkungan multikultural", "Menghormati perbedaan dalam kerja tim lintas budaya"],
    "Integritas Global": ["Menjunjung kejujuran akademik dalam tugas dan ujian", "Menggunakan teknologi digital secara bijaksana dan bertanggung jawab", "Konsisten pada nilai dan prinsip meski dalam lingkungan multikultural"],
}

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, color="1F2937", size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_doc_table(doc, data, widths):
    table = doc.add_table(rows=0, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(data):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], str(value), bold=r_idx == 0, color="FFFFFF" if r_idx == 0 else "1F2937", size=8 if r_idx == 0 else 8.2)
            set_cell_shading(cells[i], "1F4D78" if r_idx == 0 else ("F4F7FA" if r_idx % 2 == 0 else "FFFFFF"))
    return table

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)

def build_docx(path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor(31,41,55)
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [("Heading 1",16,"1F4D78",14,7),("Heading 2",12.5,"2E74B5",10,5),("Heading 3",11,"1F4D78",8,3)]:
        st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color); st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    header = sec.header.paragraphs[0]; header.text = "BM400 | Penjelasan Framework Tambahan Tiga Pilar"; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor(107,114,128)
    footer = sec.footer.paragraphs[0]; footer.text = "Dokumen penjelasan untuk stakeholder | SMA Bakti Mulya 400"; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8); footer.runs[0].font.color.rgb = RGBColor(107,114,128)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("PENJELASAN FRAMEWORK TAMBAHAN BM400"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(11,37,69)
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
    r=p.add_run("Konversi Naskah Akademis Tiga Pilar menjadi kriteria CMS dan basis pengetahuan AI"); r.font.size=Pt(12); r.font.color.rgb=RGBColor(75,85,99)
    add_doc_table(doc, [["Organisasi", "SMA Bakti Mulya 400 Jakarta"], ["Organization ID", ORG_ID], ["Sumber", "Naskah Akademis Tiga Pilar Pendidikan"], ["Status", "Framework tambahan khusus BM400"]], [1.45,5.25])

    doc.add_heading("Ringkasan Eksekutif", level=1)
    doc.add_paragraph("Framework BM400 bukan pengganti framework CIA/CDS yang sudah ada. Ia adalah lapisan tambahan khusus untuk SMA Bakti Mulya 400 yang menambahkan 13 dimensi dari Naskah Akademis Tiga Pilar: Religius, Nasionalis, dan Internasionalis.")
    doc.add_paragraph("Konversi dilakukan agar dokumen kebijakan yang bersifat naratif dapat dipakai secara konsisten oleh CMS: setiap dimensi dijadikan satu tema, setiap tema memiliki satu indikator operasional, dan setiap indikator memiliki 3-4 sub-indikator yang dapat diamati dalam percakapan asesmen.")
    add_doc_table(doc, [["Lapisan", "Sebelum", "Tambahan BM400", "Total BM400"], ["Karakter", "40 tema", "+5 tema", "45 tema"], ["Mental", "34 tema", "+4 tema", "38 tema"], ["Soft Skill", "14 tema", "+4 tema", "18 tema"], ["Total dimensi/tema tambahan", "-", "13", "13"]], [2.2,1.2,1.3,1.5])

    doc.add_heading("1. Prinsip Konversi", level=1)
    doc.add_paragraph("Dokumen sumber memiliki struktur pilar -> dimensi -> uraian operasional -> contoh kejadian/kompetensi. Struktur tersebut dipertahankan, lalu diterjemahkan ke struktur CMS yang sudah dipakai aplikasi.")
    add_doc_table(doc, [["Struktur sumber PDF", "Struktur CMS/AI", "Fungsi"], ["Pilar", "Kategori", "Karakter, Mental, atau Soft Skill"], ["Dimensi", "Tema", "Unit utama yang dilacak dalam perkembangan siswa"], ["Uraian dimensi", "Penjelasan tema", "Konteks makna dan alasan pentingnya tema"], ["Kompetensi/contoh kejadian", "Sub-indikator", "Bukti perilaku yang dapat ditandai terpenuhi atau menurun"]], [2.2,2.0,2.5])
    doc.add_paragraph("Pengelompokan kategori menggunakan sifat utama kompetensinya: nilai dan perilaku moral masuk Karakter; pola pikir, daya adaptasi, dan orientasi belajar masuk Mental; kemampuan berhubungan dan bekerja sama masuk Soft Skill.", style="Intense Quote")

    doc.add_heading("2. Pemetaan 13 Dimensi", level=1)
    doc.add_paragraph("Tabel berikut menunjukkan jejak konversi dari dimensi dalam dokumen sumber sampai ke tema dan indikator yang disimpan di lib/data/orgs/bm400.ts.")
    table_data = [["Pilar", "Dimensi sumber", "Kategori CMS", "Tema hasil konversi", "Indikator", "Sub" ]]
    table_data += [list(r) for r in rows]
    add_doc_table(doc, table_data, [0.85,1.15,0.8,2.05,1.75,0.4])

    doc.add_heading("3. Contoh Konversi Lengkap", level=1)
    doc.add_heading("Contoh A - Religius / Akhlak", level=2)
    doc.add_paragraph("Dimensi Akhlak tidak dimasukkan sebagai kalimat abstrak saja. Ia dioperasionalkan menjadi bukti perilaku yang dapat muncul dalam wawancara, misalnya pengembalian barang, pengakuan kesalahan, kesantunan, dan tanggung jawab.")
    add_doc_table(doc, [["Level", "Hasil"], ["Kategori", "Karakter"], ["Tema", "Akhlak mulia dalam interaksi sehari-hari"], ["Indikator", "Akhlak Islami dalam Keseharian"], ["Sub-indikator", "Mengembalikan barang yang ditemukan kepada pemiliknya; Mengakui kesalahan yang diperbuat; Berbicara santun kepada guru dan teman; Menepati janji dan menyelesaikan tugas dengan penuh tanggung jawab"]], [1.35,5.35])
    doc.add_heading("Contoh B - Internasionalis / Berwawasan Global", level=2)
    doc.add_paragraph("Dimensi Berwawasan Global ditempatkan di Mental karena fokusnya adalah keterbukaan cara pandang, pemahaman isu, dan kebiasaan mencari pengetahuan. Ia tidak dikategorikan sebagai Soft Skill meskipun dapat berdampak pada interaksi sosial.")
    add_doc_table(doc, [["Level", "Hasil"], ["Kategori", "Mental"], ["Tema", "Berwawasan global dan terbuka terhadap keberagaman"], ["Indikator", "Keterbukaan Wawasan Global"], ["Sub-indikator", "Menunjukkan keterbukaan terhadap budaya, bahasa, dan cara pandang yang berbeda; Memahami isu-isu global secara objektif; Mengikuti perkembangan ilmu pengetahuan dan isu global dari sumber terpercaya"]], [1.35,5.35])

    doc.add_heading("4. Bentuk Data di Aplikasi", level=1)
    doc.add_paragraph("Di dalam kode, setiap tema mengikuti bentuk data framework yang sama dengan CIA/CDS sehingga seluruh halaman CMS dapat memakai resolver yang sama. Contoh konseptualnya:")
    code = doc.add_paragraph(); code.style = "No Spacing"
    code.add_run('Kategori: Mental\n  Tema: Berwawasan global dan terbuka terhadap keberagaman\n    Indikator: Keterbukaan Wawasan Global\n      Sub-indikator: Memahami isu-isu global secara objektif').font.name = "Courier New"
    doc.add_paragraph("ID tema melanjutkan nomor framework dasar agar tidak berbenturan: Karakter 41-45, Mental 35-38, dan Soft Skill 15-18. Metadata tambahan juga dipakai untuk pengelompokan UI: Mental menggunakan group Wawasan Global, sedangkan Soft Skill menggunakan quality Kolaborasi.")

    doc.add_heading("5. Bagaimana Masuk ke Database dan AI", level=1)
    add_doc_table(doc, [["Komponen", "Yang disimpan/dilakukan", "Tujuan"], ["cia_criteria", "Satu baris per sub-indikator; organization_id BM400; search_text; embedding", "RAG kriteria penilaian"], ["pdf_knowledge", "29 chunk pengetahuan dari halaman 3-32; organization_id BM400; embedding", "RAG referensi naratif dan konteks"], ["RPC pencarian", "Mengembalikan baris global + baris yang organization_id-nya sesuai BM400", "Tenant lain tidak melihat materi BM400"], ["Rescoring", "Baris org-specific mendapat boost +0.12", "Kriteria BM400 lebih mudah muncul saat relevan"], ["Post-processing", "Nama tema/indikator AI dicocokkan lagi ke framework BM400", "Data tersimpan memakai nama kanonik dan tidak hilang"]], [1.25,3.35,2.1])
    doc.add_paragraph("Penting: embedding tidak mengubah isi kriteria. Embedding hanya menjadi representasi matematis dari search_text agar sistem dapat mencari kriteria yang paling relevan dengan percakapan. Nama tema, indikator, dan sub-indikator yang disimpan tetap berasal dari framework kanonik.", style="Intense Quote")

    doc.add_heading("6. Contoh Uji yang Mudah Dipahami", level=1)
    add_bullets(doc, [
        "Akhlak dan Integritas Global: siswa menemukan dompet lalu mengembalikannya, mengakui pernah menyontek, dan belajar menggunakan teknologi secara bertanggung jawab.",
        "Wawasan Global dan Kolaborasi Lintas Budaya: siswa bekerja dalam proyek bersama siswa dari negara lain, belajar memahami perbedaan budaya, dan membagi tugas secara adil.",
        "Demokrasi dan Toleransi: siswa menyampaikan usulan berbeda dalam rapat, menerima hasil musyawarah, dan membela teman yang diejek karena latar belakangnya.",
    ])
    doc.add_paragraph("Pada report detail, tema yang berasal dari lapisan BM400 akan diberi label BM400 - Tiga Pilar. Label ini adalah penanda asal framework untuk stakeholder, bukan nilai tambahan dan bukan kategori baru.")

    doc.add_page_break()
    doc.add_heading("7. Batasan dan Keputusan Desain", level=1)
    add_bullets(doc, [
        "Framework BM400 bersifat tambahan, sehingga seluruh tema CIA/CDS dasar tetap tersedia.",
        "Satu dimensi dibuat menjadi satu tema dan satu indikator agar struktur tetap mudah dipahami serta tidak membebani AI dengan terlalu banyak unit penilaian.",
        "Sub-indikator diringkas menjadi kalimat perilaku yang dapat diamati; ringkasan ini bukan kutipan baru yang berdiri sendiri, melainkan bentuk operasional untuk kebutuhan asesmen.",
        "RAG BM400 hanya ditampilkan untuk organization_id BM400. Organisasi lain tetap memakai kriteria dan pengetahuan global.",
        "AI tidak dipercaya untuk menentukan nama kanonik akhir; server mencocokkan hasil AI kembali ke framework lokal sebelum laporan disimpan.",
    ])

    doc.add_page_break()
    doc.add_heading("Lampiran - Rekap Sub-indikator", level=1)
    appendix = [["Dimensi", "Sub-indikator"]]
    for dim, subs in details.items(): appendix.append([dim, "\n".join("- " + s for s in subs)])
    add_doc_table(doc, appendix, [1.5,5.2])
    doc.save(path)

def build_markdown(path):
    lines = [
        "# Penjelasan Framework Tambahan BM400",
        "## Konversi Naskah Akademis Tiga Pilar menjadi CMS dan RAG AI",
        "",
        f"**Organisasi:** SMA Bakti Mulya 400 Jakarta  \n**Organization ID:** `{ORG_ID}`  \n**Sumber:** Naskah Akademis Tiga Pilar Pendidikan",
        "",
        "## Ringkasan Eksekutif",
        "Framework BM400 bukan pengganti CIA/CDS. Ini adalah lapisan tambahan khusus BM400 yang menerjemahkan 13 dimensi dari Pilar Religius, Nasionalis, dan Internasionalis menjadi kriteria CMS yang dapat diamati, dicari oleh RAG, dan disimpan di laporan.",
        "",
        "### Hasil akhir",
        "| Kategori | Framework dasar | Tambahan BM400 | Total BM400 |",
        "|---|---:|---:|---:|",
        "| Karakter | 40 tema | +5 | 45 tema |",
        "| Mental | 34 tema | +4 | 38 tema |",
        "| Soft Skill | 14 tema | +4 | 18 tema |",
        "",
        "## 1. Cara Konversi",
        "| Struktur di PDF | Struktur di CMS/AI | Penjelasan |",
        "|---|---|---|",
        "| Pilar | Kategori | Karakter, Mental, atau Soft Skill |",
        "| Dimensi | Tema | Unit utama yang dilacak dalam perkembangan siswa |",
        "| Uraian dimensi | Penjelasan tema | Konteks makna dan alasan pentingnya tema |",
        "| Kompetensi/contoh kejadian | Sub-indikator | Bukti perilaku yang dapat ditandai |",
        "",
        "Pengelompokan menggunakan sifat utama kompetensinya: nilai moral masuk **Karakter**; pola pikir, adaptasi, dan orientasi belajar masuk **Mental**; kemampuan berhubungan dan bekerja sama masuk **Soft Skill**.",
        "",
        "## 2. Pemetaan 13 Dimensi",
        "| Pilar | Dimensi | Kategori CMS | Tema hasil konversi | Indikator | Jumlah sub |",
        "|---|---|---|---|---|---:|",
    ]
    for r in rows: lines.append("| " + " | ".join(r) + " |")
    lines += [
        "",
        "## 3. Contoh Konversi",
        "### Akhlak -> Karakter",
        "Dimensi Akhlak dioperasionalkan menjadi perilaku: mengembalikan barang, mengakui kesalahan, berbicara santun, dan bertanggung jawab. Hasilnya adalah tema **Akhlak mulia dalam interaksi sehari-hari** dengan indikator **Akhlak Islami dalam Keseharian**.",
        "",
        "### Berwawasan Global -> Mental",
        "Dimensi Berwawasan Global ditempatkan di Mental karena fokusnya pada keterbukaan cara pandang, pemahaman isu, dan kebiasaan mencari pengetahuan. Hasilnya adalah tema **Berwawasan global dan terbuka terhadap keberagaman** dengan indikator **Keterbukaan Wawasan Global**.",
        "",
        "## 4. Bentuk Data",
        "```text\nKategori: Mental\n  Tema: Berwawasan global dan terbuka terhadap keberagaman\n    Indikator: Keterbukaan Wawasan Global\n      Sub-indikator: Memahami isu-isu global secara objektif\n```",
        "ID tema melanjutkan framework dasar: Karakter 41-45, Mental 35-38, Soft Skill 15-18.",
        "",
        "## 5. Masuk ke Database dan AI",
        "- `cia_criteria`: satu baris per sub-indikator, dengan `organization_id` BM400 dan embedding untuk RAG kriteria.",
        "- `pdf_knowledge`: 29 chunk pengetahuan dari halaman 3-32 PDF, dengan `organization_id` BM400 dan embedding untuk RAG referensi.",
        "- RPC pencarian: mengembalikan pengetahuan global dan pengetahuan yang khusus untuk BM400.",
        "- Rescoring: baris khusus organisasi mendapat boost `+0.12` ketika relevan.",
        "- Post-processing: hasil AI dicocokkan kembali ke framework lokal sebelum disimpan, sehingga nama kanonik tetap konsisten.",
        "",
        "## 6. Contoh Uji",
        "1. **Akhlak dan Integritas Global:** siswa menemukan dompet lalu mengembalikannya, mengakui pernah menyontek, dan memakai teknologi secara bertanggung jawab.",
        "2. **Wawasan Global dan Kolaborasi Lintas Budaya:** siswa bekerja dalam proyek dengan siswa dari negara lain, memahami perbedaan budaya, dan membagi tugas secara adil.",
        "3. **Demokrasi dan Toleransi:** siswa menyampaikan usulan berbeda dalam rapat, menerima hasil musyawarah, dan membela teman yang diejek karena latar belakangnya.",
        "",
        "Di report detail, tema BM400 diberi label **BM400 - Tiga Pilar**. Label ini hanya menunjukkan asal framework, bukan nilai tambahan atau kategori baru.",
        "",
        "## 7. Batasan Desain",
        "- BM400 adalah tambahan, bukan pengganti CIA/CDS.",
        "- Satu dimensi menjadi satu tema dan satu indikator agar mudah dipahami dan tidak membebani AI.",
        "- Sub-indikator diringkas menjadi kalimat perilaku yang dapat diamati.",
        "- RAG BM400 hanya aktif untuk organization_id BM400.",
        "- Server mencocokkan hasil AI ke framework lokal sebelum menyimpan laporan.",
        "",
        "## Lampiran: Sub-indikator",
    ]
    for dim, subs in details.items():
        lines.append(f"### {dim}")
        lines.extend(f"- {s}" for s in subs)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")

def build_pdf(path):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BMTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#0B2545"), spaceAfter=6))
    styles.add(ParagraphStyle(name="BMSub", parent=styles["Normal"], fontSize=11, leading=15, textColor=colors.HexColor("#4B5563"), spaceAfter=12))
    styles.add(ParagraphStyle(name="BMH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="BMH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#2E74B5"), spaceBefore=9, spaceAfter=4))
    styles.add(ParagraphStyle(name="BMBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=13, textColor=colors.HexColor("#1F2937"), spaceAfter=5))
    styles.add(ParagraphStyle(name="BMSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9, textColor=colors.HexColor("#1F2937")))
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=16*mm, leftMargin=16*mm, topMargin=16*mm, bottomMargin=15*mm)
    story = []
    P=lambda text, style="BMBody": Paragraph(text.replace("&", "&amp;") if "&" in text and "<" not in text else text, styles[style])
    story += [P("PENJELASAN FRAMEWORK TAMBAHAN BM400", "BMTitle"), P("Konversi Naskah Akademis Tiga Pilar menjadi kriteria CMS dan basis pengetahuan AI", "BMSub")]
    meta=[[P("Organisasi", "BMSmall"),P("SMA Bakti Mulya 400 Jakarta", "BMSmall")],[P("Organization ID", "BMSmall"),P(ORG_ID,"BMSmall")],[P("Sumber", "BMSmall"),P("Naskah Akademis Tiga Pilar Pendidikan","BMSmall")]]
    t=Table(meta,colWidths=[38*mm,132*mm]); t.setStyle(TableStyle([("BACKGROUND",(0,0),(0,-1),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#CBD5E1")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); story += [t,Spacer(1,8)]
    story += [P("Ringkasan Eksekutif","BMH1"),P("Framework BM400 bukan pengganti CIA/CDS. Ini adalah lapisan tambahan khusus BM400 yang menerjemahkan 13 dimensi dari Pilar Religius, Nasionalis, dan Internasionalis menjadi kriteria CMS yang dapat diamati, dicari oleh RAG, dan disimpan di laporan.")]
    summary=[[P(x,"BMSmall") for x in row] for row in [["Kategori","Dasar","Tambahan","Total BM400"],["Karakter","40","+5","45"],["Mental","34","+4","38"],["Soft Skill","14","+4","18"]]]
    t=Table(summary,colWidths=[55*mm,35*mm,35*mm,45*mm]); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4D78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#CBD5E1")),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F4F7FA")]),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); story += [t]
    story += [P("1. Prinsip Konversi","BMH1"),P("Struktur sumber dipertahankan lalu diterjemahkan ke struktur CMS: pilar menjadi kategori, dimensi menjadi tema, uraian menjadi penjelasan tema, dan kompetensi/contoh kejadian menjadi sub-indikator.")]
    story += [P("2. Pemetaan 13 Dimensi","BMH1")]
    mapdata=[[P(x,"BMSmall") for x in ["Pilar","Dimensi","Kategori","Tema hasil","Indikator","Sub"]]]+[[P(x,"BMSmall") for x in r] for r in rows]
    t=Table(mapdata,colWidths=[22*mm,28*mm,23*mm,48*mm,43*mm,10*mm],repeatRows=1); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4D78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.25,colors.HexColor("#CBD5E1")),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F4F7FA")]),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); story += [t]
    story += [P("3. Bentuk Data dan Ingestion","BMH1"),P("Setiap tema memakai bentuk data yang sama dengan framework CIA/CDS. ID tema melanjutkan nomor framework dasar: Karakter 41-45, Mental 35-38, dan Soft Skill 15-18.")]
    ingest=[[P(x,"BMSmall") for x in ["Komponen","Peran"]],[P("cia_criteria","BMSmall"),P("Satu baris per sub-indikator; organization_id BM400; embedding untuk RAG kriteria.","BMSmall")],[P("pdf_knowledge","BMSmall"),P("29 chunk pengetahuan dari halaman 3-32 PDF; embedding untuk RAG referensi.","BMSmall")],[P("RAG + post-processing","BMSmall"),P("Baris BM400 difilter berdasarkan organisasi, diberi boost saat relevan, lalu nama AI dicocokkan kembali ke framework kanonik.","BMSmall")]]
    t=Table(ingest,colWidths=[43*mm,128*mm]); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4D78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#CBD5E1")),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F4F7FA")]),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); story += [t]
    story += [P("4. Cara Uji","BMH1"),P("Uji dengan percakapan tentang: (1) siswa mengembalikan dompet, mengakui pernah menyontek, dan memakai teknologi secara bertanggung jawab; (2) proyek bersama siswa dari negara lain; (3) perbedaan pendapat dalam rapat dan pembelaan terhadap teman yang mengalami perundungan karena latar belakangnya."),P("Pada report detail, tema yang berasal dari lapisan BM400 akan memiliki label BM400 - Tiga Pilar. Label tersebut hanya menunjukkan asal framework, bukan nilai tambahan atau kategori baru.")]
    story += [PageBreak(), P("5. Batasan Desain","BMH1"),P("BM400 adalah tambahan, bukan pengganti CIA/CDS. Satu dimensi dibuat menjadi satu tema dan satu indikator agar mudah dipahami. Sub-indikator diringkas menjadi kalimat perilaku yang dapat diamati. RAG BM400 hanya aktif untuk organization_id BM400.")]
    story += [PageBreak(), P("Lampiran - Rekap Sub-indikator","BMH1")]
    for dim, subs in details.items():
        story += [P(dim,"BMH2"), P("; ".join(subs))]
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFont("Helvetica",7); canvas.setFillColor(colors.HexColor("#6B7280")); canvas.drawCentredString(A4[0]/2, 8*mm, f"BM400 | Penjelasan Framework Tambahan | Halaman {doc.page}"); canvas.restoreState()
    doc.build(story, onFirstPage=footer, onLaterPages=footer)

md = OUT / "penjelasan-framework-bm400.md"
docx = OUT / "penjelasan-framework-bm400.docx"
pdf = OUT / "penjelasan-framework-bm400.pdf"
build_markdown(md); build_docx(docx); build_pdf(pdf)
print(md); print(docx); print(pdf)
